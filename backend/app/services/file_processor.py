import os
import asyncio
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import pandas as pd
import json
try:
    import defusedxml.ElementTree as ET
    XML_PARSER_SAFE = True
except ImportError:
    import xml.etree.ElementTree as ET
    import warnings
    warnings.warn("defusedxml не установлен. Используется небезопасный XML парсер.", UserWarning)
    XML_PARSER_SAFE = False
from pathlib import Path
import zipfile
import rarfile
import tempfile
import shutil
from io import BytesIO, StringIO

# Импорты для работы с различными форматами
try:
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    openpyxl = None

try:
    import xlrd
    import xlwt
except ImportError:
    xlrd = None
    xlwt = None

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from sqlalchemy.orm import Session
from ..database import SessionLocal
from .. import crud, schemas

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class EnhancedFileProcessor:
    """Улучшенный процессор файлов с поддержкой 12 форматов"""
    
    SUPPORTED_FORMATS = {
        'csv': 'text/csv',
        'xls': 'application/vnd.ms-excel',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'json': 'application/json',
        'xml': 'application/xml',
        'zip': 'application/zip',
        'rar': 'application/x-rar-compressed',
        'txt': 'text/plain',
        'tsv': 'text/tab-separated-values'
    }
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        self.progress_callbacks = {}
    
    def __del__(self):
        """Очистка временных файлов"""
        try:
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except:
            pass
    
    def get_db(self) -> Session:
        """Получает сессию базы данных"""
        return SessionLocal()
    
    async def process_file(self, file_path: str, order_id: int, user_id: int) -> Dict[str, Any]:
        """Основная функция обработки файла"""
        db = self.get_db()
        try:
            # Получаем информацию о заказе
            order = crud.order_crud.get_order(db, order_id)
            if not order:
                raise ValueError(f"Заказ {order_id} не найден")
            
            # Обновляем статус заказа
            crud.update_order(db, order_id, {
                "status": "processing",
                "progress_percentage": 0.0
            })
            
            # Определяем формат файла
            file_format = self._detect_file_format(file_path)
            logger.info(f"Обработка файла {file_path} формата {file_format}")
            
            # Записываем аналитику (заглушка)
            logger.info(f"Аналитика: начало обработки файла {file_format}")
            
            # Обрабатываем файл в зависимости от формата
            result = await self._process_by_format(file_path, file_format, order_id, db)
            
            # Обновляем заказ с результатами
            crud.update_order(db, order_id, {
                "status": "completed",
                "progress_percentage": 100.0,
                "total_rows": result.get('total_rows', 0),
                "processed_rows": result.get('processed_rows', 0),
                "metadata": result.get('metadata', {})
            })
            
            # Записываем аналитику завершения (заглушка)
            logger.info(f"Аналитика: завершение обработки файла")
            
            logger.info(f"Файл {file_path} успешно обработан")
            return result
            
        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
            
            # Обновляем статус заказа на ошибку
            crud.update_order(db, order_id, {
                "status": "cancelled",
                "metadata": {"error": str(e)}
            })
            
            # Записываем аналитику ошибки (заглушка)
            logger.error(f"Аналитика: ошибка обработки файла - {str(e)}")
            
            raise
        finally:
            db.close()
    
    def _detect_file_format(self, file_path: str) -> str:
        """Определяет формат файла по расширению"""
        extension = Path(file_path).suffix.lower().lstrip('.')
        if extension in self.SUPPORTED_FORMATS:
            return extension
        raise ValueError(f"Неподдерживаемый формат файла: {extension}")
    
    async def _process_by_format(self, file_path: str, file_format: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает файл в зависимости от формата"""
        
        processors = {
            'csv': self._process_csv,
            'xls': self._process_xls,
            'xlsx': self._process_xlsx,
            'pdf': self._process_pdf,
            'doc': self._process_doc,
            'docx': self._process_docx,
            'json': self._process_json,
            'xml': self._process_xml,
            'zip': self._process_zip,
            'rar': self._process_rar,
            'txt': self._process_txt,
            'tsv': self._process_tsv
        }
        
        processor = processors.get(file_format)
        if not processor:
            raise ValueError(f"Процессор для формата {file_format} не найден")
        
        return await processor(file_path, order_id, db)
    
    async def _process_csv(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает CSV файлы"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'iso-8859-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Не удалось определить кодировку файла")
            
            return await self._process_dataframe(df, order_id, db, "CSV")
            
        except Exception as e:
            logger.error(f"Ошибка обработки CSV: {str(e)}")
            raise
    
    async def _process_xlsx(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает XLSX файлы"""
        if not openpyxl:
            raise ImportError("openpyxl не установлен")
        
        try:
            # Читаем все листы
            excel_file = pd.ExcelFile(file_path, engine='openpyxl')
            all_data = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
                df['_sheet_name'] = sheet_name
                all_data.append(df)
            
            # Объединяем все листы
            combined_df = pd.concat(all_data, ignore_index=True) if all_data else pd.DataFrame()
            
            return await self._process_dataframe(combined_df, order_id, db, "XLSX")
            
        except Exception as e:
            logger.error(f"Ошибка обработки XLSX: {str(e)}")
            raise
    
    async def _process_xls(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает XLS файлы"""
        try:
            # Читаем все листы
            excel_file = pd.ExcelFile(file_path, engine='xlrd')
            all_data = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine='xlrd')
                df['_sheet_name'] = sheet_name
                all_data.append(df)
            
            # Объединяем все листы
            combined_df = pd.concat(all_data, ignore_index=True) if all_data else pd.DataFrame()
            
            return await self._process_dataframe(combined_df, order_id, db, "XLS")
            
        except Exception as e:
            logger.error(f"Ошибка обработки XLS: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает PDF файлы"""
        if not pdfplumber:
            raise ImportError("pdfplumber не установлен")
        
        try:
            text_data = []
            table_data = []
            
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    # Обновляем прогресс
                    progress = (page_num / total_pages) * 100
                    crud.update_order(db, order_id, {
                        "progress_percentage": progress
                    })
                    
                    # Извлекаем текст
                    text = page.extract_text()
                    if text:
                        text_data.append({
                            'page': page_num + 1,
                            'text': text.strip()
                        })
                    
                    # Извлекаем таблицы
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            table_data.extend([
                                {
                                    'page': page_num + 1,
                                    'table': table_num + 1,
                                    'row': row_num,
                                    'data': row
                                }
                                for row_num, row in enumerate(table)
                            ])
            
            # Создаем изменения для текстовых данных
            changes = []
            for item in text_data:
                changes.append(schemas.OrderChangeCreate(
                    order_id=order_id,
                    row_number=item['page'],
                    column_name='text_content',
                    new_value=item['text'],
                    change_type=schemas.ChangeType.NEW
                ))
            
            # Создаем изменения для табличных данных
            for item in table_data:
                changes.append(schemas.OrderChangeCreate(
                    order_id=order_id,
                    row_number=item['row'],
                    column_name=f"page_{item['page']}_table_{item['table']}",
                    new_value=str(item['data']),
                    change_type=schemas.ChangeType.NEW
                ))
            
            # Сохраняем изменения
            if changes:
                for change_data in changes:
                    crud.create_order_change(db, change_data.dict())
            
            return {
                'total_rows': len(text_data) + len(table_data),
                'processed_rows': len(text_data) + len(table_data),
                'metadata': {
                    'pages': total_pages,
                    'text_blocks': len(text_data),
                    'tables_found': len(table_data),
                    'format': 'PDF'
                }
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {str(e)}")
            raise
    
    async def _process_docx(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает DOCX файлы"""
        if not Document:
            raise ImportError("python-docx не установлен")
        
        try:
            doc = Document(file_path)
            text_data = []
            table_data = []
            
            # Извлекаем параграфы
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_data.append({
                        'paragraph': para_num + 1,
                        'text': paragraph.text.strip()
                    })
            
            # Извлекаем таблицы
            for table_num, table in enumerate(doc.tables):
                for row_num, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append({
                        'table': table_num + 1,
                        'row': row_num + 1,
                        'data': row_data
                    })
            
            # Создаем изменения
            changes = []
            
            # Для текстовых данных
            for item in text_data:
                changes.append(schemas.OrderChangeCreate(
                    order_id=order_id,
                    row_number=item['paragraph'],
                    column_name='paragraph_text',
                    new_value=item['text'],
                    change_type=schemas.ChangeType.NEW
                ))
            
            # Для табличных данных
            for item in table_data:
                changes.append(schemas.OrderChangeCreate(
                    order_id=order_id,
                    row_number=item['row'],
                    column_name=f"table_{item['table']}",
                    new_value=str(item['data']),
                    change_type=schemas.ChangeType.NEW
                ))
            
            # Сохраняем изменения
            if changes:
                for change_data in changes:
                    crud.create_order_change(db, change_data.dict())
            
            return {
                'total_rows': len(text_data) + len(table_data),
                'processed_rows': len(text_data) + len(table_data),
                'metadata': {
                    'paragraphs': len(text_data),
                    'tables': len(doc.tables),
                    'format': 'DOCX'
                }
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки DOCX: {str(e)}")
            raise
    
    async def _process_json(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает JSON файлы"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Преобразуем JSON в плоскую структуру
            flattened_data = self._flatten_json(data)
            
            # Создаем DataFrame
            if isinstance(flattened_data, list):
                df = pd.DataFrame(flattened_data)
            else:
                df = pd.DataFrame([flattened_data])
            
            return await self._process_dataframe(df, order_id, db, "JSON")
            
        except Exception as e:
            logger.error(f"Ошибка обработки JSON: {str(e)}")
            raise
    
    async def _process_xml(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает XML файлы"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Преобразуем XML в список словарей
            xml_data = self._xml_to_dict(root)
            
            # Создаем DataFrame
            df = pd.DataFrame(xml_data)
            
            return await self._process_dataframe(df, order_id, db, "XML")
            
        except Exception as e:
            logger.error(f"Ошибка обработки XML: {str(e)}")
            raise
    
    async def _process_zip(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает ZIP архивы"""
        try:
            extracted_files = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # Извлекаем все файлы во временную директорию
                extract_dir = os.path.join(self.temp_dir, f"zip_{order_id}")
                os.makedirs(extract_dir, exist_ok=True)
                zip_ref.extractall(extract_dir)
                
                # Обрабатываем каждый извлеченный файл
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        file_path_inner = os.path.join(root, file)
                        try:
                            file_format = self._detect_file_format(file_path_inner)
                            if file_format in ['csv', 'xlsx', 'xls', 'json', 'xml', 'txt', 'tsv']:
                                result = await self._process_by_format(file_path_inner, file_format, order_id, db)
                                extracted_files.append({
                                    'filename': file,
                                    'format': file_format,
                                    'result': result
                                })
                        except Exception as e:
                            logger.warning(f"Не удалось обработать файл {file}: {str(e)}")
            
            total_rows = sum(f['result'].get('total_rows', 0) for f in extracted_files)
            processed_rows = sum(f['result'].get('processed_rows', 0) for f in extracted_files)
            
            return {
                'total_rows': total_rows,
                'processed_rows': processed_rows,
                'metadata': {
                    'extracted_files': len(extracted_files),
                    'files_processed': [f['filename'] for f in extracted_files],
                    'format': 'ZIP'
                }
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки ZIP: {str(e)}")
            raise
    
    async def _process_rar(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает RAR архивы"""
        try:
            extracted_files = []
            
            with rarfile.RarFile(file_path, 'r') as rar_ref:
                # Извлекаем все файлы во временную директорию
                extract_dir = os.path.join(self.temp_dir, f"rar_{order_id}")
                os.makedirs(extract_dir, exist_ok=True)
                rar_ref.extractall(extract_dir)
                
                # Обрабатываем каждый извлеченный файл
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        file_path_inner = os.path.join(root, file)
                        try:
                            file_format = self._detect_file_format(file_path_inner)
                            if file_format in ['csv', 'xlsx', 'xls', 'json', 'xml', 'txt', 'tsv']:
                                result = await self._process_by_format(file_path_inner, file_format, order_id, db)
                                extracted_files.append({
                                    'filename': file,
                                    'format': file_format,
                                    'result': result
                                })
                        except Exception as e:
                            logger.warning(f"Не удалось обработать файл {file}: {str(e)}")
            
            total_rows = sum(f['result'].get('total_rows', 0) for f in extracted_files)
            processed_rows = sum(f['result'].get('processed_rows', 0) for f in extracted_files)
            
            return {
                'total_rows': total_rows,
                'processed_rows': processed_rows,
                'metadata': {
                    'extracted_files': len(extracted_files),
                    'files_processed': [f['filename'] for f in extracted_files],
                    'format': 'RAR'
                }
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки RAR: {str(e)}")
            raise
    
    async def _process_txt(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает TXT файлы"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'iso-8859-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Не удалось определить кодировку файла")
            
            # Разбиваем на строки
            lines = content.split('\n')
            
            # Создаем изменения
            changes = []
            for line_num, line in enumerate(lines):
                if line.strip():
                    changes.append(schemas.OrderChangeCreate(
                        order_id=order_id,
                        row_number=line_num + 1,
                        column_name='text_line',
                        new_value=line.strip(),
                        change_type=schemas.ChangeType.NEW
                    ))
            
            # Сохраняем изменения
            if changes:
                for change_data in changes:
                    crud.create_order_change(db, change_data.dict())
            
            return {
                'total_rows': len(lines),
                'processed_rows': len([l for l in lines if l.strip()]),
                'metadata': {
                    'total_lines': len(lines),
                    'non_empty_lines': len([l for l in lines if l.strip()]),
                    'format': 'TXT'
                }
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки TXT: {str(e)}")
            raise
    
    async def _process_tsv(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает TSV файлы"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'iso-8859-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, sep='\t', encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Не удалось определить кодировку файла")
            
            return await self._process_dataframe(df, order_id, db, "TSV")
            
        except Exception as e:
            logger.error(f"Ошибка обработки TSV: {str(e)}")
            raise
    
    async def _process_doc(self, file_path: str, order_id: int, db: Session) -> Dict[str, Any]:
        """Обрабатывает DOC файлы (заглушка)"""
        # DOC файлы требуют специальных библиотек или конвертации
        # Для простоты возвращаем базовую информацию
        return {
            'total_rows': 1,
            'processed_rows': 1,
            'metadata': {
                'format': 'DOC',
                'note': 'DOC файлы требуют дополнительной обработки'
            }
        }
    
    async def _process_dataframe(self, df: pd.DataFrame, order_id: int, db: Session, format_name: str) -> Dict[str, Any]:
        """Обрабатывает DataFrame и создает изменения"""
        try:
            total_rows = len(df)
            changes = []
            
            # Обрабатываем каждую строку
            for idx, (row_idx, row) in enumerate(df.iterrows()):
                # Обновляем прогресс
                if idx % 100 == 0:  # Обновляем каждые 100 строк
                    progress = (idx / total_rows) * 100
                    crud.update_order(db, order_id, {
                        "progress_percentage": progress
                    })
                
                # Создаем изменения для каждой колонки
                for col_name, value in row.items():
                    if pd.notna(value):  # Пропускаем NaN значения
                        change_data = {
                            "order_id": order_id,
                            "row_number": idx + 1,
                            "column_name": str(col_name),
                            "new_value": str(value),
                            "change_type": "new"
                        }
                        crud.create_order_change(db, change_data)
            
            return {
                'total_rows': total_rows,
                'processed_rows': total_rows,
                'metadata': {
                    'columns': list(df.columns),
                    'column_count': len(df.columns),
                    'format': format_name
                }
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки DataFrame: {str(e)}")
            raise
    
    def _flatten_json(self, data: Any, parent_key: str = '', sep: str = '.') -> Dict[str, Any]:
        """Преобразует вложенный JSON в плоскую структуру"""
        items = []
        
        if isinstance(data, dict):
            for k, v in data.items():
                new_key = f"{parent_key}{sep}{k}" if parent_key else k
                if isinstance(v, (dict, list)):
                    items.extend(self._flatten_json(v, new_key, sep=sep).items())
                else:
                    items.append((new_key, v))
        elif isinstance(data, list):
            for i, v in enumerate(data):
                new_key = f"{parent_key}{sep}{i}" if parent_key else str(i)
                if isinstance(v, (dict, list)):
                    items.extend(self._flatten_json(v, new_key, sep=sep).items())
                else:
                    items.append((new_key, v))
        else:
            return {parent_key: data}
        
        return dict(items)
    
    def _xml_to_dict(self, element) -> List[Dict[str, Any]]:
        """Преобразует XML элемент в список словарей"""
        result = []
        
        def parse_element(elem, path=""):
            current_path = f"{path}.{elem.tag}" if path else elem.tag
            
            # Если элемент имеет текст
            if elem.text and elem.text.strip():
                result.append({
                    'path': current_path,
                    'value': elem.text.strip(),
                    'attributes': elem.attrib
                })
            
            # Обрабатываем атрибуты
            for attr_name, attr_value in elem.attrib.items():
                result.append({
                    'path': f"{current_path}@{attr_name}",
                    'value': attr_value,
                    'attributes': {}
                })
            
            # Рекурсивно обрабатываем дочерние элементы
            for child in elem:
                parse_element(child, current_path)
        
        parse_element(element)
        return result
    
    async def save_to_storage(self, filename: str, content: bytes) -> str:
        """Сохраняет файл в облачное хранилище"""
        try:
            # Генерируем уникальное имя файла
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_filename = f"{timestamp}_{filename}"
            
            # В реальном приложении здесь будет загрузка в DigitalOcean Spaces
            # Пока сохраняем локально
            storage_dir = os.path.join(os.getcwd(), 'backend', 'uploads')
            os.makedirs(storage_dir, exist_ok=True)
            
            storage_path = os.path.join(storage_dir, unique_filename)
            
            with open(storage_path, 'wb') as f:
                f.write(content)
            
            # Возвращаем относительный путь
            return f"uploads/{unique_filename}"
            
        except Exception as e:
            logger.error(f"Ошибка сохранения файла в хранилище: {str(e)}")
            raise
    
    def validate_file(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Валидирует файл"""
        try:
            # Определяем расширение файла
            extension = Path(filename).suffix.lower().lstrip('.')
            
            # Проверяем поддерживаемые форматы
            if extension not in self.SUPPORTED_FORMATS:
                return {
                    'valid': False,
                    'errors': [f'Неподдерживаемый формат файла: {extension}'],
                    'file_type': extension
                }
            
            # Проверяем размер файла (максимум 100MB)
            if len(content) > 100 * 1024 * 1024:
                return {
                    'valid': False,
                    'errors': ['Файл слишком большой (максимум 100MB)'],
                    'file_type': extension
                }
            
            # Проверяем, что файл не пустой
            if len(content) == 0:
                return {
                    'valid': False,
                    'errors': ['Файл пустой'],
                    'file_type': extension
                }
            
            # Дополнительные проверки для конкретных форматов
            validation_errors = []
            
            if extension in ['csv', 'tsv']:
                # Проверяем CSV/TSV файлы
                try:
                    # Пробуем прочитать как текст
                    text_content = content.decode('utf-8')
                    if not text_content.strip():
                        validation_errors.append('CSV/TSV файл не содержит данных')
                except UnicodeDecodeError:
                    try:
                        # Пробуем другие кодировки
                        content.decode('cp1251')
                    except UnicodeDecodeError:
                        validation_errors.append('Не удалось определить кодировку файла')
            
            elif extension == 'json':
                # Проверяем JSON файлы
                try:
                    json.loads(content.decode('utf-8'))
                except (json.JSONDecodeError, UnicodeDecodeError):
                    validation_errors.append('Некорректный JSON формат')
            
            elif extension == 'xml':
                # Проверяем XML файлы
                try:
                    ET.fromstring(content)
                except ET.ParseError:
                    validation_errors.append('Некорректный XML формат')
            
            elif extension in ['xlsx', 'xls']:
                # Проверяем Excel файлы
                try:
                    # Создаем временный файл для проверки
                    temp_file = os.path.join(self.temp_dir, f"temp_{filename}")
                    with open(temp_file, 'wb') as f:
                        f.write(content)
                    
                    if extension == 'xlsx' and openpyxl:
                        pd.read_excel(temp_file, engine='openpyxl')
                    elif extension == 'xls':
                        pd.read_excel(temp_file, engine='xlrd')
                    
                    # Удаляем временный файл
                    os.remove(temp_file)
                    
                except Exception:
                    validation_errors.append(f'Некорректный {extension.upper()} файл')
            
            elif extension == 'pdf':
                # Проверяем PDF файлы
                if not content.startswith(b'%PDF'):
                    validation_errors.append('Некорректный PDF файл')
            
            elif extension in ['zip', 'rar']:
                # Проверяем архивы
                try:
                    temp_file = os.path.join(self.temp_dir, f"temp_{filename}")
                    with open(temp_file, 'wb') as f:
                        f.write(content)
                    
                    if extension == 'zip':
                        with zipfile.ZipFile(temp_file, 'r') as zip_ref:
                            zip_ref.testzip()
                    elif extension == 'rar':
                        with rarfile.RarFile(temp_file, 'r') as rar_ref:
                            rar_ref.testrar()
                    
                    os.remove(temp_file)
                    
                except Exception:
                    validation_errors.append(f'Поврежденный {extension.upper()} архив')
            
            # Возвращаем результат валидации
            return {
                'valid': len(validation_errors) == 0,
                'errors': validation_errors,
                'file_type': extension,
                'file_size': len(content),
                'mime_type': self.SUPPORTED_FORMATS.get(extension, 'application/octet-stream')
            }
            
        except Exception as e:
            logger.error(f"Ошибка валидации файла: {str(e)}")
            return {
                'valid': False,
                'errors': [f'Ошибка валидации: {str(e)}'],
                'file_type': extension if 'extension' in locals() else 'unknown'
            }

# Глобальный экземпляр процессора
file_processor = EnhancedFileProcessor()

# Функция для использования в других модулях
async def process_file_async(file_path: str, order_id: int, user_id: int) -> Dict[str, Any]:
    """Асинхронная обработка файла"""
    return await file_processor.process_file(file_path, order_id, user_id)
